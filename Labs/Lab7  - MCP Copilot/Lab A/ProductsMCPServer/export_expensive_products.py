import asyncio
import httpx
from typing import List

from StoreUtils import Product

API_URL = "https://fakestoreapi.com/products"
OUTPUT_DOCX = "expensive_products.docx"
OUTPUT_TXT = "expensive_products.txt"


async def fetch_products() -> List[Product]:
    async with httpx.AsyncClient() as client:
        resp = await client.get(API_URL)
        resp.raise_for_status()
        data = resp.json()
        return Product.list_from_api(data)


def filter_expensive(products: List[Product], threshold: float = 50.0) -> List[Product]:
    return [p for p in products if p.price > threshold]


def save_to_docx(products: List[Product], path: str) -> None:
    try:
        from docx import Document
    except Exception:
        raise

    doc = Document()
    doc.add_heading("Expensive Products (price > 50)", level=1)
    for p in products:
        doc.add_heading(p.title, level=2)
        doc.add_paragraph(p.description)
    doc.save(path)


def save_to_txt(products: List[Product], path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write("Expensive Products (price > 50)\n\n")
        for p in products:
            f.write(p.title + "\n")
            f.write(p.description + "\n\n")


async def main():
    products = await fetch_products()
    expensive = filter_expensive(products, 50.0)

    if not expensive:
        print("No products over 50 found.")
        return

    # Try saving as docx first
    try:
        save_to_docx(expensive, OUTPUT_DOCX)
        print(f"Saved {len(expensive)} products to {OUTPUT_DOCX}")
    except Exception:
        save_to_txt(expensive, OUTPUT_TXT)
        print(f"docx not available; saved {len(expensive)} products to {OUTPUT_TXT}")


if __name__ == "__main__":
    asyncio.run(main())
